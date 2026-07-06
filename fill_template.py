"""
Fill the internship report docx template with content.
Reads content from docs/企业实习中期报告_填充内容.md and fills into the template.
"""
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TEMPLATE_PATH = "docs/2023090912011_谢沁桐_软件工程（数字信息处理）_企业实习中期报告模板.docx"
OUTPUT_PATH = "docs/企业实习中期报告_已填充.docx"
FILL_MD_PATH = "docs/企业实习中期报告_填充内容.md"


# ── Step 1: Parse the markdown fill-in document into sections ──────────

def parse_fill_md(path):
    """Parse markdown fill document into a dict: section_key -> list of paragraph texts."""
    with open(path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    sections = {}
    current_key = None
    current_lines = []
    in_code_block = False

    section_map = {
        "### 1.1": "1.1", "### 1.2": "1.2", "### 1.3": "1.3", "### 1.4": "1.4",
        "## 2": "2",
        "## 3": "3_intro",
        "### 3.1": "3.1", "### 3.2": "3.2", "### 3.3": "3.3", "### 3.4": "3.4",
        "## 4": "4_intro",
        "### 4.1": "4.1", "### 4.2": "4.2",
        "## 5": "5_intro",
        "### 5.1": "5.1", "### 5.2": "5.2",
        "## 参考文献": "refs",
    }

    content_start = 0
    for i, line in enumerate(lines):
        if line.strip().startswith("## 1 毕业实习"):
            content_start = i
            break

    for line in lines[content_start:]:
        stripped = line.strip()

        if stripped.startswith("```"):
            in_code_block = not in_code_block
            continue
        if in_code_block:
            continue

        matched_key = None
        for marker, key in section_map.items():
            if stripped.startswith(marker):
                matched_key = key
                break

        if matched_key:
            if current_key and current_lines:
                sections[current_key] = _clean_paragraphs(current_lines)
            current_key = matched_key
            current_lines = []
            continue

        if current_key and stripped:
            current_lines.append(stripped)

    if current_key and current_lines:
        sections[current_key] = _clean_paragraphs(current_lines)

    return sections


def _clean_paragraphs(lines):
    """Clean markdown formatting, split into paragraphs at natural boundaries."""
    # First pass: clean each line, join consecutive non-blank lines
    cleaned_lines = []
    for line in lines:
        # Skip table rows, separators, image refs
        if line.startswith("|") or line.startswith("---") or line.startswith("!["):
            if cleaned_lines and cleaned_lines[-1] != "":
                cleaned_lines.append("")  # mark paragraph break
            continue
        cleaned = line
        cleaned = re.sub(r'\*\*([^*]+)\*\*', r'\1', cleaned)
        cleaned = re.sub(r'\*([^*]+)\*', r'\1', cleaned)
        cleaned = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', cleaned)
        cleaned = re.sub(r'^[-*]\s+', '', cleaned)
        cleaned = re.sub(r'^#{1,6}\s+', '', cleaned)  # heading markers
        cleaned = cleaned.strip()
        cleaned_lines.append(cleaned)

    # Second pass: group into paragraphs
    # Paragraph break on: empty lines, sub-headings, reference entries
    paragraphs = []
    current = []
    for line in cleaned_lines:
        if not line:
            if current:
                paragraphs.append(" ".join(current))
                current = []
            continue

        # Reference entries: each [N] starts a new paragraph
        if re.match(r'^\[\d+\]', line):
            if current:
                paragraphs.append(" ".join(current))
                current = []
            current.append(line)
            paragraphs.append(" ".join(current))
            current = []
            continue

        # Sub-headings with colons start new paragraph
        if (line.endswith("：") or line.endswith(":")) and len(line) < 80:
            if current:
                paragraphs.append(" ".join(current))
                current = []
            current.append(line)
            continue

        # Chinese numbered sub-sections start new paragraph
        if re.match(r'^（[一二三四五六七八九十]+）', line):
            if current:
                paragraphs.append(" ".join(current))
                current = []
            current.append(line)
            continue

        # Regular content
        current.append(line)

    if current:
        paragraphs.append(" ".join(current))

    # Filter very short paragraphs, merge with previous if needed
    result = []
    for p in paragraphs:
        p = p.strip()
        if not p:
            continue
        if len(p) < 10 and result:
            result[-1] = result[-1] + " " + p
        else:
            result.append(p)

    return result


# ── Step 2: Fill the template ──────────────────────────────────────────

def add_paragraph_after(doc, paragraph, text, font_size=Pt(12), font_name="宋体"):
    """Insert a new paragraph after the given paragraph."""
    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    # Wrap as Paragraph - get the document part
    from docx.text.paragraph import Paragraph
    new_para = Paragraph(new_p, paragraph._parent)

    # Add text run with formatting
    pPr = new_p.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        new_p.insert(0, pPr)
    # Set font
    run_elem = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rPr.append(rFonts)
    sz = OxmlElement("w:sz")
    sz.set(qn('w:val'), str(int(font_size.pt * 2)))
    rPr.append(sz)
    run_elem.append(rPr)
    t = OxmlElement("w:t")
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    run_elem.append(t)
    new_p.append(run_elem)

    return new_para


def set_paragraph_text(paragraph, text, font_size=Pt(12), font_name="宋体"):
    """Replace paragraph text with formatted text."""
    # Remove all existing runs
    for r in paragraph._element.findall(qn('w:r')):
        paragraph._element.remove(r)

    # Add new run
    run = paragraph.add_run(text)
    run.font.size = font_size
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def fill_section(doc, para_idx, content_paragraphs, font_size=Pt(12)):
    """Replace instruction paragraph at para_idx with multi-paragraph content."""
    if not content_paragraphs:
        return

    instr_para = doc.paragraphs[para_idx]
    set_paragraph_text(instr_para, content_paragraphs[0], font_size=font_size)

    insert_point = instr_para
    for text in content_paragraphs[1:]:
        insert_point = add_paragraph_after(doc, insert_point, text, font_size=font_size)


def delete_paragraph(paragraph):
    """Delete a paragraph from the document."""
    p = paragraph._element
    p.getparent().remove(p)


def main():
    print("Parsing fill content from markdown...")
    sections = parse_fill_md(FILL_MD_PATH)
    print(f"Found {len(sections)} sections")
    for k, v in sections.items():
        print(f"  {k}: {len(v)} paragraphs")

    print("\nLoading template...")
    doc = Document(TEMPLATE_PATH)
    print(f"Template has {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables")

    # ── Mapping: section_key -> instruction_paragraph_index ──
    section_map = {
        "1.1": 30, "1.2": 32, "1.3": 34, "1.4": 36,
        "2": 51,
        "3_intro": 81,
        "3.1": 83, "3.2": 85, "3.3": 87, "3.4": 89,
        "4.1": 106, "4.2": 108,
        "5.1": 128, "5.2": 130,
    }

    # ── Process in REVERSE order ──
    sorted_sections = sorted(section_map.items(), key=lambda x: x[1], reverse=True)

    for key, para_idx in sorted_sections:
        if key in sections:
            print(f"  Filling section {key} at P{para_idx}: {len(sections[key])} paragraphs")
            fill_section(doc, para_idx, sections[key])
        elif key == "3_intro":
            # No intro content between ## 3 and ### 3.1 — clear the instruction
            print(f"  Clearing section 3 intro at P{para_idx}")
            set_paragraph_text(doc.paragraphs[para_idx],
                "本人在系统中主要负责数据影像采集子系统的设计与实现，以下为各模块的详细方案说明。",
                font_size=Pt(12))

    # ── Remove instruction note P90 ──
    print("\nRemoving instruction notes...")
    for i in range(len(doc.paragraphs) - 1, -1, -1):
        text = doc.paragraphs[i].text.strip()
        # Remove notes that start with 注：or 注意：
        if text.startswith("注：") or text.startswith("注意："):
            delete_paragraph(doc.paragraphs[i])
            print(f"  Deleted note at P{i}")

    # ── Remove 说明 section P24-P26 ──
    for i in range(26, 23, -1):
        if i < len(doc.paragraphs):
            text = doc.paragraphs[i].text.strip()
            if any(kw in text for kw in ["说明", "6000", "模板为参考"]):
                delete_paragraph(doc.paragraphs[i])
                print(f"  Deleted instruction at P{i}")

    # ── Update references ──
    print("\nUpdating references...")
    refs_content = sections.get("refs", [])
    if refs_content:
        ref_start = None
        for i, p in enumerate(doc.paragraphs):
            if p.style.name == "Heading 1" and "参考" in p.text:
                ref_start = i
                break

        if ref_start:
            old_refs = []
            for i in range(ref_start + 1, min(ref_start + 6, len(doc.paragraphs))):
                if doc.paragraphs[i].text.strip():
                    old_refs.append(i)
            for i in reversed(old_refs):
                delete_paragraph(doc.paragraphs[i])

            ref_para = None
            for i, p in enumerate(doc.paragraphs):
                if p.style.name == "Heading 1" and "参考" in p.text:
                    ref_para = p
                    break

            if ref_para:
                insert_point = ref_para
                for ref_text in refs_content:
                    if ref_text.startswith("[") and len(ref_text) > 5:
                        insert_point = add_paragraph_after(doc, insert_point, ref_text,
                                                           font_size=Pt(10.5))
                print(f"  Inserted {len(refs_content)} references")

    # ── Fill table 5-1 ──
    print("\nFilling table 5-1...")
    table_data = [
        ["1", "影像报告模块完整开发（文件上传、报告检索、文件管理）", "第7周周一", "第7周周五"],
        ["2", "数据统计汇总接口完善（多维度统计、趋势分析数据接口）", "第7周周一", "第7周周三"],
        ["3", "前后端接口联调（数据采集模块全部接口）", "第7周周四", "第8周周三"],
        ["4", "前端数据采集页面开发（列表页、详情页、新增/编辑表单）", "第8周周一", "第8周周五"],
        ["5", "集成测试（数据采集模块与其他模块的协作验证）", "第9周周一", "第9周周三"],
        ["6", "系统测试与缺陷修复", "第9周周四", "第9周周五"],
        ["7", "系统部署（测试环境部署与验证）", "第10周周一", "第10周周二"],
        ["8", "文档完善（用户手册中数据采集模块部分）", "第10周周二", "第10周周三"],
        ["9", "项目答辩准备（演示环境搭建、PPT制作）", "第10周周三", "第10周周五"],
    ]

    if doc.tables:
        table = doc.tables[0]
        # Remove all rows except header
        tbl = table._tbl
        while len(table.rows) > 1:
            row = table.rows[1]
            tbl.remove(row._tr)

        # Add data rows
        for row_data in table_data:
            row = table.add_row()
            for j, cell_text in enumerate(row_data):
                if j < len(row.cells):
                    cell = row.cells[j]
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.text = ""
                    run = cell.paragraphs[0].add_run(cell_text)
                    run.font.size = Pt(10.5)
                    run.font.name = "宋体"
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), "宋体")
        print(f"  Table filled with {len(table_data)} rows")

    # ── Remove leading empty paragraphs that are now excessive ──
    # Clean up excessive empty paragraphs between major sections
    print("\nCleaning up...")
    empty_count = 0
    for i in range(len(doc.paragraphs) - 1, -1, -1):
        p = doc.paragraphs[i]
        if not p.text.strip() and p.style.name == "Normal":
            # Don't delete if it's between content (keep as spacing)
            empty_count += 1
    print(f"  Found {empty_count} empty paragraphs (kept for spacing)")

    # ── Save ──
    print(f"\nSaving to {OUTPUT_PATH}...")
    doc.save(OUTPUT_PATH)
    print("Done! Output: " + OUTPUT_PATH)


if __name__ == "__main__":
    main()
