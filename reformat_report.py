"""
Reformat the filled internship report according to formatting standards.
Reference: docs/附件2：实习报告的书写规范.docx
Input: docs/企业实习中期报告_已填充 - 副本.docx
"""
import re
from docx import Document
from docx.shared import Pt, Cm, Emu, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


INPUT_PATH = "docs/企业实习中期报告_已填充 - 副本.docx"

# ── Helpers ────────────────────────────────────────────────────────────

def set_run_font(run, cn_font="宋体", en_font="Times New Roman", size=Pt(12), bold=None):
    """Set font for a run: Chinese font + English font + size."""
    run.font.name = en_font
    run.font.size = size
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), cn_font)
    rFonts.set(qn('w:ascii'), en_font)
    rFonts.set(qn('w:hAnsi'), en_font)
    rFonts.set(qn('w:cs'), en_font)
    if bold is not None:
        run.font.bold = bold


def set_paragraph_spacing(para, line_spacing=None, space_before=None, space_after=None,
                          first_line_indent=None):
    """Set paragraph spacing properties."""
    pf = para.paragraph_format
    if line_spacing is not None:
        pf.line_spacing = line_spacing
    if space_before is not None:
        pf.space_before = space_before
    if space_after is not None:
        pf.space_after = space_after
    if first_line_indent is not None:
        pf.first_line_indent = first_line_indent


def set_paragraph_font(para, cn_font="宋体", en_font="Times New Roman",
                       size=Pt(12), bold=None):
    """Set font for all runs in a paragraph."""
    for run in para.runs:
        if run.text.strip():
            set_run_font(run, cn_font, en_font, size, bold)


def ensure_paragraph_has_run(para, text=None):
    """Ensure paragraph has at least one run with text."""
    if not para.runs:
        run = para.add_run(text or "")
        return run
    return para.runs[0]


# ── Main ───────────────────────────────────────────────────────────────

def main():
    print("Loading document...")
    doc = Document(INPUT_PATH)
    print(f"  {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables")

    # ── 1. Fix body paragraphs (Normal style, NOT headings, NOT refs) ──
    print("\n1. Formatting body paragraphs...")
    body_count = 0
    in_refs = False
    for i, p in enumerate(doc.paragraphs):
        style_name = p.style.name if p.style else ""
        text = p.text.strip()

        # Track if we're in the references section
        if 'Heading' in style_name and "参考" in text:
            in_refs = True
            continue
        if in_refs and 'Heading' in style_name:
            in_refs = False

        # Skip TOC paragraphs, headings, references, empty
        if 'toc' in style_name.lower():
            continue
        if 'Heading' in style_name:
            continue
        if in_refs:
            continue
        if not text:
            continue

        # Skip cover page content (first ~22 paragraphs)
        if i < 22:
            continue

        # This is a body paragraph
        set_paragraph_font(p, cn_font="宋体", en_font="Times New Roman", size=Pt(12))
        set_paragraph_spacing(p, line_spacing=1.5, first_line_indent=Cm(0.74))
        body_count += 1

    print(f"  Formatted {body_count} body paragraphs")

    # ── 2. Fix headings ──
    print("\n2. Formatting headings...")
    heading_counts = {"Heading 1": 0, "Heading 2": 0, "Heading 3": 0}
    for i, p in enumerate(doc.paragraphs):
        style_name = p.style.name if p.style else ""
        text = p.text.strip()
        if not text:
            continue

        if style_name == "Heading 1":
            set_paragraph_font(p, cn_font="黑体", en_font="Times New Roman", size=Pt(16), bold=True)
            pf = p.paragraph_format
            pf.line_spacing = 1.5
            pf.space_before = Pt(12)
            pf.space_after = Pt(6)
            pf.first_line_indent = Cm(0)
            heading_counts["Heading 1"] += 1

        elif style_name == "Heading 2":
            set_paragraph_font(p, cn_font="黑体", en_font="Times New Roman", size=Pt(14), bold=True)
            pf = p.paragraph_format
            pf.line_spacing = 1.5
            pf.space_before = Pt(6)
            pf.space_after = Pt(3)
            pf.first_line_indent = Cm(0)
            heading_counts["Heading 2"] += 1

        elif style_name == "Heading 3":
            set_paragraph_font(p, cn_font="黑体", en_font="Times New Roman", size=Pt(12), bold=True)
            pf = p.paragraph_format
            pf.line_spacing = 1.5
            pf.space_before = Pt(3)
            pf.space_after = Pt(3)
            pf.first_line_indent = Cm(0)
            heading_counts["Heading 3"] += 1

    for k, v in heading_counts.items():
        print(f"  {k}: {v}")

    # ── 3. Fix table formatting ──
    print("\n3. Formatting tables...")
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        set_run_font(run, cn_font="宋体", en_font="Times New Roman",
                                     size=Pt(10.5))
        print(f"  Table formatted: {len(table.rows)} rows x {len(table.columns)} cols")

    # ── 4. Fix TOC formatting ──
    print("\n4. Formatting TOC...")
    toc_count = 0
    for p in doc.paragraphs:
        if 'toc' in (p.style.name or "").lower():
            for run in p.runs:
                set_run_font(run, cn_font="宋体", en_font="Times New Roman", size=Pt(12))
            toc_count += 1
    print(f"  Formatted {toc_count} TOC entries")

    # ── 5. Set TOC title to 黑体 ──
    for p in doc.paragraphs:
        if 'toc' in (p.style.name or "").lower() and p.text.strip() == "目 录":
            set_paragraph_font(p, cn_font="黑体", en_font="Times New Roman", size=Pt(16), bold=True)
            set_paragraph_spacing(p, line_spacing=1.5)
            print("  TOC title set to 黑体")
            break

    # ── 6. Fix reference section ──
    print("\n5. Formatting references...")
    ref_count = 0
    in_refs = False
    for p in doc.paragraphs:
        style_name = p.style.name or ""
        text = p.text.strip()

        if style_name == "Heading 1" and "参考" in text:
            in_refs = True
            continue
        if in_refs and 'Heading' in style_name:
            in_refs = False
            continue
        if in_refs and text:
            set_paragraph_font(p, cn_font="宋体", en_font="Times New Roman", size=Pt(10.5))
            # Fixed 20pt line spacing, no indent
            pf = p.paragraph_format
            pf.line_spacing = Pt(20)
            pf.first_line_indent = Cm(0)
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
            ref_count += 1
    print(f"  Formatted {ref_count} references")

    # ── 7. Fix cover page header ──
    # First few paragraphs should be bold with 1.5 line spacing
    print("\n6. Formatting cover page...")
    for i in range(min(22, len(doc.paragraphs))):
        p = doc.paragraphs[i]
        if p.text.strip():
            set_paragraph_font(p, cn_font="宋体", en_font="Times New Roman", size=Pt(12))
            set_paragraph_spacing(p, line_spacing=1.5)

    # ── 8. Ensure all numbers and English use Times New Roman ──
    print("\n7. Checking English/number fonts...")
    # This is handled by the body paragraph formatting above

    # ── Save ──
    print(f"\nSaving reformatted document to {INPUT_PATH}...")
    doc.save(INPUT_PATH)
    print("Done!")


if __name__ == "__main__":
    main()
